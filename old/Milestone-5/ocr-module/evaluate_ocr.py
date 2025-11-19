#!/usr/bin/env python3
"""OCR evaluation template for Milestone 5 using EasyOCR."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, Iterable, Tuple, List

import easyocr
from pdf2image import convert_from_path
from PIL import Image
import docx

# initialize reader once (keep languages configurable)
READER = easyocr.Reader(["en"], gpu=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Milestone 5 OCR evaluation harness.")
    parser.add_argument("--input-dir", required=True, help="Folder containing images/PDF/DOCX assets.")
    parser.add_argument("--ground-truth", help="Optional path to JSON with reference text.")
    parser.add_argument("--out", default="ocr_eval_results.json", help="Destination JSON report.")
    parser.add_argument("--max-samples", type=int, default=0, help="Optional limit for quick smoke tests.")
    return parser.parse_args()


# ---------- Helper: Edit distance for CER/WER ----------
def edit_distance(a: List[str], b: List[str]) -> int:
    """Classic Levenshtein edit distance on lists of tokens (chars or words)."""
    n, m = len(a), len(b)
    dp = [[0] * (m + 1) for _ in range(n + 1)]
    for i in range(n + 1):
        dp[i][0] = i
    for j in range(m + 1):
        dp[0][j] = j
    for i in range(1, n + 1):
        for j in range(1, m + 1):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            dp[i][j] = min(
                dp[i - 1][j] + 1,
                dp[i][j - 1] + 1,
                dp[i - 1][j - 1] + cost,
            )
    return dp[n][m]

# ---------- 1. Text presence detection ----------
def detect_text_presence(file_path: Path) -> Tuple[bool, float]:
    """
    Detect text using EasyOCR.
    ANY non-empty OCR result means text is present 
    (confidence may be low, but text EXISTS).
    """
    suffix = file_path.suffix.lower()

    # -------- DOCX detection --------
    if suffix == ".docx":
        try:
            doc = docx.Document(str(file_path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text)
            has = len(text.strip()) > 0
            return has, 1.0 if has else 0.0
        except Exception:
            return False, 0.0

    # -------- PDF first page --------
    if suffix == ".pdf":
        try:
            pages = convert_from_path(str(file_path), dpi=200)
            if not pages:
                return False, 0.0
            img = pages[0]   # PIL Image
        except Exception:
            return False, 0.0
    else:
        # -------- Image file --------
        try:
            img = Image.open(str(file_path)).convert("RGB")
        except Exception:
            return False, 0.0

    # -------- Convert PIL to numpy (CRITICAL FIX) --------
    import numpy as np
    np_img = np.array(img)

    # -------- OCR detection --------
    try:
        results = READER.readtext(np_img, detail=1)
    except Exception as e:
        print("OCR ERROR:", e)
        return False, 0.0

    if not results:
        return False, 0.0

    confidences = [float(item[2]) for item in results if len(item) >= 3]
    avg_conf = sum(confidences)/len(confidences) if confidences else 0.0

    return True, avg_conf




# ---------- 2. Extract text ----------
def extract_text(file_path: Path) -> str:
    """
    Extract text from images, PDFs, and DOCX files using EasyOCR.
    Returns a clean string with all detected lines.
    """

    suffix = file_path.suffix.lower()

    # -------- DOCX --------
    if suffix == ".docx":
        try:
            doc = docx.Document(str(file_path))
            return "\n".join(
                p.text for p in doc.paragraphs if p.text and p.text.strip()
            )
        except Exception:
            return ""

    texts: List[str] = []

    # -------- PDF --------
    if suffix == ".pdf":
        try:
            pages = convert_from_path(str(file_path), dpi=200)
        except Exception:
            return ""

        import numpy as np

        for page in pages:
            try:
                np_img = np.array(page)        # PIL -> NumPy
                result = READER.readtext(np_img, detail=1)
            except Exception as e:
                print("PDF OCR error:", e)
                continue

            for item in result:
                if len(item) >= 2 and item[1]:
                    texts.append(item[1])

        return "\n".join(texts)

    # -------- IMAGES --------
    try:
        img = Image.open(str(file_path)).convert("RGB")
        import numpy as np
        np_img = np.array(img)
        result = READER.readtext(np_img, detail=1)
    except Exception:
        return ""

    for item in result:
        if len(item) >= 2 and item[1]:
            texts.append(item[1])

    return "\n".join(texts)




# ---------- 3. Quality metrics ----------
def compute_quality_metrics(pred_text: str, ref_text: str | None) -> Dict[str, Any]:
    metrics: Dict[str, Any] = {}
    if ref_text is None:
        return metrics

    # Character error rate (CER)
    pred_chars = list(pred_text)
    ref_chars = list(ref_text)
    cer = edit_distance(pred_chars, ref_chars) / max(1, len(ref_chars))

    # Word error rate (WER)
    pred_words = pred_text.split()
    ref_words = ref_text.split()
    wer = edit_distance(pred_words, ref_words) / max(1, len(ref_words))

    metrics["char_error_rate"] = cer
    metrics["word_error_rate"] = wer
    return metrics


# ---------- Utility: load references ----------
def load_references(path: str | None) -> Dict[str, str]:
    if not path:
        return {}
    path_obj = Path(path)
    if path_obj.suffix.lower() == ".json":
        try:
            data = json.loads(path_obj.read_text(encoding="utf-8"))
            return {item["id"]: item["text"] for item in data}
        except Exception:
            raise RuntimeError("Ground truth JSON must be list of {'id','text'} items.")
    else:
        raise NotImplementedError("Only JSON reference files supported.")


# ---------- Document iterator ----------
def iter_documents(input_dir: Path) -> Iterable[Path]:
    for ext in ("*.png", "*.jpg", "*.jpeg", "*.tiff", "*.pdf", "*.docx"):
        yield from input_dir.rglob(ext)


# ---------- Main ----------
def main() -> None:
    args = parse_args()
    input_dir = Path(args.input_dir)
    reference_map = load_references(args.ground_truth)

    records: list[Dict[str, Any]] = []
    for idx, path in enumerate(iter_documents(input_dir)):
        if args.max_samples and idx >= args.max_samples:
            break

        has_text, confidence = detect_text_presence(path)
        extracted_text = extract_text(path) if has_text else ""
        ref_text = reference_map.get(path.stem)
        metrics = compute_quality_metrics(extracted_text, ref_text)

        records.append(
            {
                "file": str(path),
                "has_text": has_text,
                "confidence": confidence,
                "extracted_text": extracted_text if len(extracted_text) <= 256 else (extracted_text[:256] + "..."),
                "metrics": metrics,
            }
        )

    report = {"input_dir": str(input_dir), "samples": len(records), "results": records}
    Path(args.out).write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(f"Wrote OCR evaluation report to {args.out}")


if __name__ == "__main__":
    main()
